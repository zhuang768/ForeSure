import os
import logging
from datetime import datetime
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.opc.constants import RELATIONSHIP_TYPE
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import RGBColor

logger = logging.getLogger(__name__)

# The report is bilingual: every heading and label is printed as "中文 / English", and every section written by
# the agent is printed twice, Chinese first, then the English version the agent produced in the "<field>_en" key.
_ENGLISH_GREY = RGBColor(0x55, 0x55, 0x55)
_MISSING_EN = "（未提供英文版 / English version not provided）"
_MISSING_TRANSLATION = "（未提供翻譯 / translation not provided）"
_LINK_BLUE = "0563C1"
_ENGLISH_NAME_MAX_CHARS = 60  # keeps file names within filesystem limits for long English product names


def _heading(doc, zh: str, en: str, level: int):
    return doc.add_heading(f"{zh} / {en}", level)


def _label(zh: str, en: str, value) -> str:
    return f"{zh} / {en}：{value}"


def _english(paragraph, text: str):
    run = paragraph.add_run(text)
    run.italic = True
    run.font.color.rgb = _ENGLISH_GREY
    return run


def _add_hyperlink(paragraph, url: str, text: str) -> None:
    """python-docx has no hyperlink API, so build the w:hyperlink element and its relationship by hand."""
    r_id = paragraph.part.relate_to(url, RELATIONSHIP_TYPE.HYPERLINK, is_external=True)
    hyperlink = OxmlElement("w:hyperlink")
    hyperlink.set(qn("r:id"), r_id)
    run = OxmlElement("w:r")
    props = OxmlElement("w:rPr")
    color = OxmlElement("w:color")
    color.set(qn("w:val"), _LINK_BLUE)
    underline = OxmlElement("w:u")
    underline.set(qn("w:val"), "single")
    props.append(color)
    props.append(underline)
    run.append(props)
    text_el = OxmlElement("w:t")
    text_el.text = text
    text_el.set(qn("xml:space"), "preserve")
    run.append(text_el)
    hyperlink.append(run)
    paragraph._p.append(hyperlink)


def _translated_pair(original: str, zh, en) -> tuple[str, str | None]:
    """(Chinese, English) for a news field whose original may be in either language and may lack a translation."""
    zh_text = zh or original
    en_text = en or (original if zh else None)
    return zh_text, en_text


def _add_trigger_news(doc, proposal_data: dict) -> None:
    """Headline and summary in both languages, with the headline and the URL as clickable links to the source."""
    link = proposal_data.get("news_link") or ""
    title_zh, title_en = _translated_pair(proposal_data.get("source_news") or "N/A",
                                          proposal_data.get("source_news_zh"), proposal_data.get("source_news_en"))
    p = doc.add_paragraph("新聞標題 / Headline：")
    if link:
        _add_hyperlink(p, link, title_zh)
    else:
        p.add_run(title_zh)
    _english(doc.add_paragraph(), title_en or _MISSING_TRANSLATION)

    summary_zh, summary_en = _translated_pair(proposal_data.get("news_summary") or "N/A",
                                              proposal_data.get("news_summary_zh"), proposal_data.get("news_summary_en"))
    doc.add_paragraph(_label("新聞摘要", "Summary", summary_zh))
    _english(doc.add_paragraph(), summary_en or _MISSING_TRANSLATION)

    if link:
        _add_hyperlink(doc.add_paragraph("新聞連結 / Source link："), link, link)


def _bilingual_section(doc, zh_text, en_text) -> None:
    """Chinese body, then the English version in italic grey; flag it honestly when the agent gave none."""
    doc.add_paragraph(zh_text or "N/A")
    _english(doc.add_paragraph(), en_text or _MISSING_EN)


def _add_actuarial_basis(doc, basis: dict) -> None:
    """Spell out where each actuarial figure comes from, separating real statistics from assumptions."""
    _heading(doc, "數據依據", "Basis", level=3)
    source = basis.get("probability_source")
    if source and source != "assumption":
        if basis.get("probability_source_en"):
            source = f"{source}（{basis['probability_source_en']}）"
        doc.add_paragraph(_label("發生機率依據", "Probability source", source)
                          + f"；{_label('方法', 'Method', basis.get('probability_method', ''))}")
    else:
        doc.add_paragraph(_label("發生機率依據", "Probability source",
                                 f"假設值，無官方統計 / assumption, no official statistics（{basis.get('probability_method', '')}）"))
    if basis.get("low_sample"):
        doc.add_paragraph(_label("注意", "Note",
                                 "嚴重事件樣本少於 5 筆，機率估計的不確定性高。"
                                 " / Fewer than 5 severe events in the sample; the probability estimate is highly uncertain."))
    loss_line = _label("單次損失依據", "Loss basis", f"假設值 / assumption（{basis.get('loss_method', '')}）")
    if basis.get("assumed_loss_per_household_usd"):
        loss_line += (f"；每戶假設損失 / assumed loss per household USD {basis['assumed_loss_per_household_usd']}"
                      f"（{basis.get('assumed_loss_note', '')}）")
    doc.add_paragraph(loss_line)
    doc.add_paragraph(_label("保費計算", "Premium method", basis.get("premium_method", "")))

    mc = basis.get("monte_carlo_gpu")
    if mc:
        _heading(doc, "AMD ROCm GPU 百萬次巨災壓力測試", "AMD ROCm Catastrophe Stress Testing", level=3)
        doc.add_paragraph(
            _label("運算引擎", "Compute Engine", f"{mc.get('engine', 'AMD ROCm GPU Tensor Core')} ({mc.get('iterations', 1000000):,} runs)")
        )
        doc.add_paragraph(
            _label("99.5% 巨災極值損失 (VaR 99.5%)", "Catastrophe VaR 99.5%", f"USD {mc.get('var_99_5_usd', 0):,.2f}（200 年一遇極端情境）")
        )
        doc.add_paragraph(
            _label("99.5% 尾端期望損失 (TVaR 99.5%)", "Tail Value at Risk 99.5%", f"USD {mc.get('tvar_99_5_usd', 0):,.2f}")
        )
        doc.add_paragraph(
            _label("資本適足性清償要求 (SCR)", "Solvency Capital Requirement", f"USD {mc.get('solvency_capital_requirement_usd', 0):,.2f}（{mc.get('capital_adequacy_status', '100% Solvency Compliant')}）")
        )
        doc.add_paragraph(
            _label("數學校準加成倍數", "Calibrated Loading", f"{mc.get('calibrated_markup_multiplier', 'N/A')}x（經證明足以吸收 99.5% 巨災暴險）")
        )


_GROUNDING_LABEL = {
    "pass": ("通過：風險與定價數字皆可追溯到來源", "Pass: every risk and pricing figure traces back to a source"),
    "warn": ("警示：需補充揭露，建議人工確認", "Warning: a disclosure is missing; human confirmation advised"),
    "fail": ("未通過：發現無來源的數字或無法驗證的引用，須人工審核後才可送審",
             "Fail: ungrounded figures or unverifiable citations found; human review required before filing"),
}
_FLAG_LABEL = {
    "unsupported_number": ("無來源的數字", "Ungrounded number"),
    "unverified_citation": ("無法驗證的引用", "Unverifiable citation"),
    "missing_disclosure": ("假設值未揭露", "Assumption not disclosed"),
}


def _add_grounding_section(doc, grounding: dict) -> None:
    """Rule-based hallucination check: the verdict, the counts and every flag, written for the human reviewer."""
    _heading(doc, "4. 幻覺檢測", "Grounding Check", level=1)
    zh, en = _GROUNDING_LABEL.get(grounding.get("status"), ("未知", "Unknown"))
    doc.add_paragraph(_label("結果", "Verdict", zh))
    _english(doc.add_paragraph(), en)
    checked, grounded = grounding.get("checked_claims", 0), grounding.get("grounded_claims", 0)
    doc.add_paragraph(_label("檢查的數字", "Figures checked",
                             f"{checked} 項，其中 {grounded} 項有來源 / {grounded} of {checked} grounded"))
    for flag in grounding.get("flags") or []:
        fz, fe = _FLAG_LABEL.get(flag.get("type"), (str(flag.get("type", "")), str(flag.get("type", ""))))
        line = f"[{flag.get('severity', '')}] {fz} / {fe}：{flag.get('field', '')}"
        if flag.get("value"):
            line += f" → {flag['value']}"
        doc.add_paragraph(line)
        if flag.get("excerpt"):
            doc.add_paragraph(f"「{flag['excerpt']}」")
        if flag.get("message"):
            doc.add_paragraph(flag["message"])
    doc.add_paragraph(_label("檢查器版本", "Checker version", grounding.get("checker_version", "")))


def _safe_name(name: str) -> str:
    return "".join(c for c in (name or "") if c.isalnum() or c == " ").strip().replace(" ", "_")


def _report_filename(proposal: dict) -> str:
    """<timestamp>_<Chinese name>[_<English name>].docx; the English part is skipped when it is already in the name."""
    zh = _safe_name(proposal.get("product_name", "Report")) or "Product_Proposal"
    parts = [datetime.now().strftime("%Y%m%d_%H%M%S"), zh]
    en = _safe_name(proposal.get("product_name_en", ""))[:_ENGLISH_NAME_MAX_CHARS].rstrip("_")
    if en and en not in zh:
        parts.append(en)
    return "_".join(parts) + ".docx"


def generate_report(proposal_data: dict, output_dir: str = "reports") -> str:
    """
    將 Agent 生成的提案資料排版並輸出為中英雙語的 Word (.docx) 檔案。
    """
    logger.info("生成 Word 提案報告中...")

    if not os.path.exists(output_dir):
        os.makedirs(output_dir)

    doc = Document()

    # 標題 / Title
    title = doc.add_heading('創新保險商品開發提案書', 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    subtitle = doc.add_paragraph('Innovative Insurance Product Proposal', style='Subtitle')
    subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER

    # 日期 / Date
    date_str = datetime.now().strftime("%Y-%m-%d %H:%M:%S")
    date_p = doc.add_paragraph(_label("生成日期", "Generated", date_str))
    date_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT

    if "error" in proposal_data:
        doc.add_paragraph("發生錯誤，無法生成完整報告。 / An error occurred and the full report could not be generated.")
        doc.add_paragraph(proposal_data["error"])
        filepath = os.path.join(output_dir, "error_report.docx")
        doc.save(filepath)
        return filepath

    proposal = proposal_data["proposal"]
    actuarial = proposal_data["actuarial_data"]

    # 1. 商品概要 / Product overview
    _heading(doc, '1. 商品名稱與市場缺口', 'Product Name & Market Gap', level=1)
    name = proposal.get('product_name', '未命名商品')
    if proposal.get('product_name_en'):
        name = f"{name} / {proposal['product_name_en']}"
    doc.add_paragraph(_label("商品名稱", "Product name", name)).style = 'Heading 3'

    _heading(doc, '觸發時事', 'Trigger Event', level=2)
    _add_trigger_news(doc, proposal_data)

    _heading(doc, '市場缺口分析', 'Market Gap Analysis', level=2)
    _bilingual_section(doc, proposal.get('market_gap'), proposal.get('market_gap_en'))

    # 2. 客群與保障 / Audience and coverage
    _heading(doc, '2. 目標客群與保障範圍', 'Target Audience & Coverage', level=1)
    _heading(doc, '目標客群', 'Target Audience', level=2)
    _bilingual_section(doc, proposal.get('target_audience'), proposal.get('target_audience_en'))

    _heading(doc, '保障細節', 'Coverage Details', level=2)
    _bilingual_section(doc, proposal.get('coverage_details'), proposal.get('coverage_details_en'))

    _heading(doc, '除外不保事項', 'Exclusions', level=2)
    _bilingual_section(doc, proposal.get('exclusions'), proposal.get('exclusions_en'))

    # 3. 精算數據與商業邏輯 / Actuarial figures and business logic
    _heading(doc, '3. 基礎精算與商業評估', 'Actuarial Basis & Business Assessment', level=1)
    _heading(doc, 'AI 初步精算估算', 'Preliminary AI Actuarial Estimate', level=2)
    p_act = doc.add_paragraph()
    p_act.add_run(_label("預估風險發生機率", "Estimated probability", f"{actuarial.get('probability_pct', '0')}%") + "\n")
    p_act.add_run(_label("單次事故預期損失", "Expected loss per event", f"USD {actuarial.get('expected_loss_usd', '0')}") + "\n")
    p_act.add_run(_label("建議保費定價區間", "Suggested premium range",
                         f"USD {actuarial['premium_range_usd'][0]} ~ USD {actuarial['premium_range_usd'][1]}"))

    basis = actuarial.get("basis")
    if basis:
        _add_actuarial_basis(doc, basis)

    _heading(doc, '商業邏輯與獲利模式', 'Business Logic & Profit Model', level=2)
    _bilingual_section(doc, proposal.get('business_logic'), proposal.get('business_logic_en'))

    # 4. 幻覺檢測 / Grounding check（舊紀錄沒有，略過）
    grounding = proposal_data.get("grounding")
    if grounding:
        _add_grounding_section(doc, grounding)

    # 儲存檔案 / Save
    filepath = os.path.join(output_dir, _report_filename(proposal))
    doc.save(filepath)

    # 上鏈存證與 audit log 寫入由 main.run_pipeline 負責（見 chain_writer / run_store）。
    logger.info(f"報告已生成並儲存至：{filepath}")
    return filepath
